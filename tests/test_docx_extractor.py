"""
test_docx_extractor.py — DOCX 분할/추출 단위 테스트

대상:
  split_docx_to_blocks (6 케이스):
    1. 빈 docx → 블록 0개, 에러 없음
    2. 표만 있는 docx (표 3개) → 블록 3개, 각 table_markdown 보유
    3. Heading 스타일 없는 docx → 폰트 휴리스틱(R2)으로 블록 분할
    4. vMerge 병합 셀 → markdown 빈 칸, merge_info 기록
    5. TOC 패턴(탭+숫자$) → F1 필터로 "(목차)" 블록 1개로 압축
    6. 단락 31개 섹션 → F3 강제 컷, "(계속)" suffix

  _extract_docx_eval_from_table (4 케이스):
    7. 정상 배점표 (단일 행 3개) → total_points=100, categories 3개
    8. vMerge 병합 배점 (배치+공간이 40 공유) → shared_with 배열 생성
    9. 소계/합계 행 → 자동 제외
   10. 배점 컬럼 없는 표 → evaluation_categories=[], total_points=None, 에러 없음

LLM/네트워크 의존 없음. 픽스처는 python-docx 인메모리 생성.

실행:
  cd <repo-root>
  backend/venv/Scripts/python.exe -m pytest tests/test_docx_extractor.py -v
"""

from __future__ import annotations

import sys
import types
from pathlib import Path

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

# ── sys.path 설정 ─────────────────────────────────────────────────────────────
_BACKEND = Path(__file__).resolve().parents[1] / "backend"
if str(_BACKEND) not in sys.path:
    sys.path.insert(0, str(_BACKEND))

# ── 무거운 의존성 스텁 (LLM/httpx 미설치 환경) ──────────────────────────────
try:
    import httpx  # noqa: F401
except ImportError:
    sys.modules["httpx"] = types.ModuleType("httpx")

if "config" not in sys.modules:
    class _Settings:
        model_id = "stub"
        model_id_classify = "stub"
        dpi_classify = 72
        dpi_extract = 120
        extraction_priority_limit = 2

    _cfg = types.ModuleType("config")
    _cfg.settings = _Settings()
    _cfg.axes_keys_for = lambda *a, **k: []
    _cfg.PAGE_TYPES = {
        "COVER", "TOC_HERO", "SITE_CONTEXT", "CONCEPT", "SPECIAL_SPACE",
        "RENDERING_EXT", "RENDERING_INT", "SITE_PLAN", "LANDSCAPE",
        "FLOOR_PLAN", "SECTION", "ELEVATION", "CIRCULATION", "HEALTH_CENTER",
        "TECHNICAL", "AREA_TABLE", "SUSTAINABILITY", "UNIT_PLAN",
        "INCENTIVE_TABLE", "BRANDING", "BUSINESS_VIABILITY", "AREA_INCREASE",
        "VIEW_ANALYSIS", "COMMUNITY_PROGRAM", "COMPANY_PORTFOLIO",
        "CONSTRUCTION_PLAN", "UNIT_PLAN_PENTHOUSE",
    }
    _cfg.BRIEF_PAGE_TYPES = {
        "BRIEF_OVERVIEW", "BRIEF_PROJECT_INFO", "BRIEF_SITE", "BRIEF_PROGRAM",
        "BRIEF_DESIGN_MASSING", "BRIEF_DESIGN_FACADE", "BRIEF_DESIGN_SUSTAIN",
        "BRIEF_DESIGN_SPECIAL", "BRIEF_DESIGN_GUIDE", "BRIEF_TECHNICAL",
        "BRIEF_REGULATIONS", "BRIEF_EVALUATION", "BRIEF_SUBMISSION", "BRIEF_ADMIN",
    }
    sys.modules["config"] = _cfg

if "services.llm_client" not in sys.modules:
    _llm = types.ModuleType("services.llm_client")
    _llm.call_messages = lambda **k: "{}"
    sys.modules["services.llm_client"] = _llm

if "services.utils" not in sys.modules:
    _utils = types.ModuleType("services.utils")
    for _n in ("get_page_text", "normalize_design_guidelines_grouped", "ocr_page",
               "parse_json_response", "rasterize_pdf", "rasterize_page_tiled",
               "safe_encode_image"):
        setattr(_utils, _n, lambda *a, **k: None)

    def _first(data, key):
        v = (data or {}).get(key) or {}
        if isinstance(v, list):
            v = v[0] if v else {}
        return v if isinstance(v, dict) else {}

    def _as_list(data, key):
        v = (data or {}).get(key) or []
        return v if isinstance(v, list) else ([v] if v else [])

    _utils._first = _first
    _utils._as_list = _as_list
    sys.modules["services.utils"] = _utils

from services.docx_loader import split_docx_to_blocks          # noqa: E402
from services.data_extractor import _extract_docx_eval_from_table  # noqa: E402


# ── 픽스처 헬퍼 ──────────────────────────────────────────────────────────────

def _save(doc: Document, tmp_path: Path, name: str = "test.docx") -> str:
    p = tmp_path / name
    doc.save(str(p))
    return str(p)


def _heading_run(doc: Document, text: str, size_pt: float = 18.0):
    """폰트 크기만으로 visual heading 만들기 (Heading 스타일 미사용)."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(size_pt)
    return para


def _add_xml_vmerge(cell, *, restart: bool):
    """셀 tcPr에 w:vMerge 요소를 추가.
    restart=True → <w:vMerge w:val="restart"/>, False → <w:vMerge/> (continue).
    """
    tcPr = cell._tc.get_or_add_tcPr()
    vmerge = OxmlElement("w:vMerge")
    if restart:
        vmerge.set(qn("w:val"), "restart")
    tcPr.append(vmerge)


def _make_eval_block(table_md: str, merge_info: list[dict] | None = None) -> dict:
    return {
        "block_num": 1,
        "header_text": "심사기준",
        "paragraphs": [],
        "table_markdown": table_md,
        "merge_info": merge_info or [],
    }


# ════════════════════════════════════════════════════════════════════════════
# split_docx_to_blocks
# ════════════════════════════════════════════════════════════════════════════

def test_empty_docx_returns_no_blocks(tmp_path):
    """빈 docx → 블록 0개, 에러 없음."""
    path = _save(Document(), tmp_path, "empty.docx")
    assert split_docx_to_blocks(path) == []


def test_tables_only_three_blocks(tmp_path):
    """표만 있는 docx (표 3개) → 블록 3개, 각 table_markdown 보유."""
    doc = Document()
    for i in range(3):
        tbl = doc.add_table(rows=2, cols=2)
        tbl.rows[0].cells[0].text = f"H1_{i}"
        tbl.rows[0].cells[1].text = f"H2_{i}"
        tbl.rows[1].cells[0].text = f"V1_{i}"
        tbl.rows[1].cells[1].text = f"V2_{i}"

    path = _save(doc, tmp_path, "tables_only.docx")
    blocks = split_docx_to_blocks(path)

    assert len(blocks) == 3
    for i, blk in enumerate(blocks):
        assert blk["table_markdown"] is not None, f"블록 {i} table_markdown 없음"
        assert f"H1_{i}" in blk["table_markdown"]
        assert f"V2_{i}" in blk["table_markdown"]


def test_no_heading_styles_splits_by_font_heuristic(tmp_path):
    """Heading 스타일 없는 docx → 폰트 휴리스틱(R2)으로 섹션 분할.

    16pt 이상 단락은 굵기 무관 헤딩으로 인식(R2: FONT_HEADING_PT_BIG=16).
    """
    doc = Document()
    _heading_run(doc, "제1장 사업개요", size_pt=16.0)
    doc.add_paragraph("본문 단락 A.")
    doc.add_paragraph("본문 단락 B.")
    _heading_run(doc, "제2장 입찰 절차", size_pt=16.0)
    doc.add_paragraph("입찰 본문.")

    path = _save(doc, tmp_path, "no_heading_styles.docx")
    blocks = split_docx_to_blocks(path)

    assert len(blocks) >= 2
    headers = [b["header_text"] for b in blocks]
    assert any("제1장 사업개요" in h for h in headers)
    assert any("제2장 입찰 절차" in h for h in headers)


def test_vmerge_blank_cells_and_merge_info_recorded(tmp_path):
    """vMerge 병합 셀 → markdown 병합 행은 빈 칸, merge_info에 병합 메타 기록.

    표 구조 (3행 × 2열):
      row0: GROUP (restart) | Detail A
      row1: (continue)      | Detail B
      row2: (continue)      | Detail C
    """
    doc = Document()
    tbl = doc.add_table(rows=3, cols=2)
    tbl.rows[0].cells[0].text = "GROUP"
    tbl.rows[0].cells[1].text = "Detail A"
    tbl.rows[1].cells[1].text = "Detail B"
    tbl.rows[2].cells[1].text = "Detail C"
    _add_xml_vmerge(tbl.rows[0].cells[0], restart=True)
    _add_xml_vmerge(tbl.rows[1].cells[0], restart=False)
    _add_xml_vmerge(tbl.rows[2].cells[0], restart=False)

    path = _save(doc, tmp_path, "vmerge.docx")
    blocks = split_docx_to_blocks(path)

    assert len(blocks) == 1
    blk = blocks[0]
    md = blk["table_markdown"]
    assert md is not None

    # markdown 행 분리: [헤더, 구분선, body row1, body row2]
    lines = md.split("\n")
    assert len(lines) >= 4
    # body row1 (rows[1]): col0 = vMerge continue → 빈 칸
    first_body = lines[2]
    assert first_body.startswith("|  |") or "|  |" in first_body

    # merge_info: GROUP 이 3행 병합
    assert len(blk["merge_info"]) >= 1
    assert any(
        m["value"] == "GROUP" and m["merged_rows"] == 3
        for m in blk["merge_info"]
    )


def test_toc_paragraphs_compressed_to_single_block(tmp_path):
    r"""F1: 탭+페이지번호(\t\d+$) 단락들 → "(목차)" 단일 블록으로 압축."""
    doc = Document()
    # 비-TOC 헤더 (visual heading)
    _heading_run(doc, "제안 요청서", size_pt=20.0)
    # TOC 항목 5개
    for i in range(1, 6):
        doc.add_paragraph(f"제{i}장 내용\t{i}")
    # 이후 일반 본문
    doc.add_paragraph("본문 시작입니다.")

    path = _save(doc, tmp_path, "toc.docx")
    blocks = split_docx_to_blocks(path)

    toc_blocks = [b for b in blocks if "(목차)" in b["header_text"]]
    assert len(toc_blocks) == 1, f"(목차) 블록 없음 — blocks={[b['header_text'] for b in blocks]}"
    # 5개 TOC 단락이 하나의 블록으로 압축
    assert len(toc_blocks[0]["paragraphs"]) == 5


def test_force_cut_31_paragraphs(tmp_path):
    """F3: 단락 31개 섹션 (Heading 1개 + 본문 30개) → 강제 컷 발생, '(계속)' suffix.

    _FORCE_CUT_PARAS = 30 이므로 블록 내 단락 수가 30에 도달하는 순간 분할.
    """
    doc = Document()
    # visual heading 1개 (섹션 시작)
    _heading_run(doc, "긴 섹션 제목", size_pt=20.0)
    # 본문 30개 → heading 포함 총 31개
    for i in range(30):
        doc.add_paragraph(f"본문 단락 {i + 1}번.")

    path = _save(doc, tmp_path, "force_cut.docx")
    blocks = split_docx_to_blocks(path)

    continues = [b for b in blocks if "(계속)" in b["header_text"]]
    assert len(continues) >= 1, (
        f"'(계속)' 블록 없음 (F3 미발동) — blocks={[b['header_text'] for b in blocks]}"
    )


# ════════════════════════════════════════════════════════════════════════════
# _extract_docx_eval_from_table
# ════════════════════════════════════════════════════════════════════════════

def test_eval_normal_table_total_100():
    """정상 배점표 (3개 단일 행) → total_points=100, 카테고리 3개."""
    md = "\n".join([
        "| 구분 | 평가사항 | 배점 |",
        "| --- | --- | --- |",
        "| 배치계획 | 배치의 적정성 | 30 |",
        "| 디자인 | 창의성 및 공공성 | 40 |",
        "| 친환경 | 인증 여부 | 30 |",
    ])
    result = _extract_docx_eval_from_table(_make_eval_block(md))

    assert result["total_points"] == 100
    cats = result["evaluation_categories"]
    assert len(cats) == 3
    names = [c["name"] for c in cats]
    assert "배치계획" in names
    assert "디자인" in names
    assert "친환경" in names
    pts = sorted(c["points"] for c in cats)
    assert pts == [30, 30, 40]


def test_eval_merged_points_col_shared_with():
    """배점 셀이 두 구분에 걸쳐 vMerge (배치+공간 40 공유) → shared_with 배열 생성.

    원본 표:
      | 구분  | 평가사항 | 배점 |
      | 배치  | x       | 40   |  ← 배점 셀 vMerge restart (row 1)
      | 공간  | y       |      |  ← 배점 셀 vMerge continue (row 2)
      | 기술  | z       | 60   |

    merge_info: 원본 table.rows 인덱스 기준 row=1 (헤더=row0, 첫 body=row1).
    """
    md = "\n".join([
        "| 구분 | 평가사항 | 배점 |",
        "| --- | --- | --- |",
        "| 배치 | x | 40 |",
        "| 공간 | y |  |",
        "| 기술 | z | 60 |",
    ])
    merge_info = [{"row": 1, "col": 2, "merged_rows": 2, "value": "40"}]
    result = _extract_docx_eval_from_table(_make_eval_block(md, merge_info))

    cats = result["evaluation_categories"]
    # 배치(shared)+공간 그룹 1개 + 기술 단일 1개 = 2개
    assert len(cats) == 2

    shared_cat = next((c for c in cats if c["shared_with"]), None)
    assert shared_cat is not None, "shared_with 카테고리 없음"
    assert shared_cat["name"] == "배치"
    assert "공간" in shared_cat["shared_with"]
    assert shared_cat["points"] == 40

    single_cat = next((c for c in cats if not c["shared_with"]), None)
    assert single_cat is not None
    assert single_cat["name"] == "기술"
    assert single_cat["points"] == 60

    assert result["total_points"] == 100


def test_eval_subtotal_rows_auto_excluded():
    """소계/합계 행 → 카테고리 목록에서 자동 제외, 명시적 합계는 total_points로 사용."""
    md = "\n".join([
        "| 구분 | 평가사항 | 배점 |",
        "| --- | --- | --- |",
        "| 배치 | x | 30 |",
        "| 디자인 | y | 40 |",
        "| 기술 | z | 30 |",
        "| 소계 |  | 70 |",
        "| 가격 | a | 30 |",
        "| 총계 |  | 100 |",
    ])
    result = _extract_docx_eval_from_table(_make_eval_block(md))

    names = [c["name"] for c in result["evaluation_categories"]]
    assert "소계" not in names, "소계 행이 카테고리로 포함됨"
    assert "총계" not in names, "총계 행이 카테고리로 포함됨"
    for expected in ("배치", "디자인", "기술", "가격"):
        assert expected in names

    # 총계 행 100이 explicit_total로 설정됨
    assert result["total_points"] == 100


def test_eval_no_points_column_returns_empty():
    """배점 컬럼 없는 표 → evaluation_categories=[], total_points=None, 에러 없음."""
    md = "\n".join([
        "| 항목 | 설명 |",
        "| --- | --- |",
        "| 사업명 | 대전 ABC 사업 |",
        "| 위치 | 대전광역시 서구 |",
    ])
    result = _extract_docx_eval_from_table(_make_eval_block(md))

    assert result["evaluation_categories"] == []
    assert result["total_points"] is None
