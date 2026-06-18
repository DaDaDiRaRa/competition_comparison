"""
docx_loader.py — DOCX 지침서를 블록 단위로 분할/파싱

설계 결정 (D-2, KT 실측 보강):
  - 블록 분할: R1~R5 + F1~F3 (TOC 필터·orphan 헤더 forward merge·force-cut)
  - page_header_text 폴백 순서: A(Heading) → B(폰트 visual) → C(캡션) → D(표 첫 행) → E(첫 단락) → F(디폴트)
  - 표: 파이프 마크다운 변환. vMerge 셀은 빈 칸 출력. 가로 병합 셀은 같은 텍스트 반복.

PDF 흐름과 독립적인 모듈 — 기존 page_classifier / data_extractor 는 수정하지 않음.
"""

from __future__ import annotations

import re
from typing import Any


# ── 설정 상수 ──────────────────────────────────────────────────────────────────
_FONT_HEADING_PT     = 14.0     # R2: 굵게+이 크기 이상
_FONT_HEADING_PT_BIG = 16.0     # R2: 굵기 무관 이 크기 이상
_HEADING_MAX_LEN     = 60       # R2: 헤딩 단락 최대 길이
_FORCE_CUT_PARAS     = 30       # F3: 단락 수
_FORCE_CUT_CHARS     = 8000     # F3: 글자 수
_HEADER_FALLBACK_LEN = 60       # page_header_text 폴백 최대 길이
_SOURCE_TEXT_CAP     = 6000     # get_block_source_text 최대 길이
_SOURCE_HEAD_CAP     = 4000     # 위 초과 시 앞쪽 보존
_SOURCE_TAIL_CAP     = 2000     # 위 초과 시 뒤쪽 보존
_TABLE_PREVIEW_ROWS  = 10       # source_text에 포함할 표 미리보기 행 수
_TABLE_CELL_MAX      = 60       # 표 셀 텍스트 최대 길이 (LLM 입력용)

# 정규식
_RE_SECTION_NUM = re.compile(r'^(제\d+장|\d+(\.\d+){0,3})\s')
_RE_CAPTION    = re.compile(r'\[\s*(표|양식|서식|별표)\s*\d+')
_RE_TOC        = re.compile(r'\t\d+\s*$')          # F1: 탭+페이지번호
_RE_FIGURE     = re.compile(r'^(그림\s*\d+|Fig\.)')
_RE_CONTINUE   = re.compile(r'\(\s*계속\s*\)')

# 설계지침 글머리 라벨 패턴 (BRIEF_DESIGN_* 계층 보존용)
# LLM 프롬프트에 "이 페이지에 N)/가)/① 등 글머리 보입니다" 힌트로 전달돼
# design_guidelines_grouped[] 추출 정확도 향상.
_LABEL_PATTERNS = {
    "1)":   re.compile(r'^\d+\)\s'),                # 1) 토지이용 (그룹 헤더)
    "가)":  re.compile(r'^[가-힣]\)\s'),             # 가) 대지 주변 (세부 항목)
    "①":    re.compile(r'^[①-⑳]\s?'),               # ① 구 청 / ① 비서실 (시설/세부 라벨)
    "·":    re.compile(r'^[•·]\s'),       # • 또는 · 글머리
    "-":    re.compile(r'^[-–—]\s'),      # - 또는 – (en/em dash)
    "Ÿ":    re.compile(r'^Ÿ\s'),                    # 영등포 PDF 의 워드 글머리 변환 잔여
    "I.":   re.compile(r'^[IVX]+\.\s'),             # I./II./III. (로마숫자 챕터)
}


# ── 폰트 휴리스틱 ───────────────────────────────────────────────────────────────
def _para_visual_heading(p) -> bool:
    """폰트 휴리스틱(R2): 굵게 AND 14pt 이상 OR 16pt 이상. 단락 길이 < 60자."""
    text = (p.text or "").strip()
    if not text or len(text) >= _HEADING_MAX_LEN:
        return False
    # 단락 내 첫 run의 폰트 정보로 판단 (정확하지는 않지만 안전한 근사)
    bold = False
    size_pt = None
    for run in p.runs:
        if run.bold:
            bold = True
        sz = getattr(run.font, "size", None)
        if sz is not None:
            # docx의 Length 단위는 EMU. .pt 속성으로 포인트 추출 가능.
            try:
                size_pt = max(size_pt or 0.0, float(sz.pt))
            except Exception:
                pass
        if bold and size_pt and size_pt >= _FONT_HEADING_PT:
            return True
    if size_pt and size_pt >= _FONT_HEADING_PT_BIG:
        return True
    return False


def _para_heading_style(p) -> bool:
    """R1: 스타일 이름이 Heading 1/2/3 또는 한국어 "제목 *"."""
    name = (getattr(p.style, "name", None) or "")
    if name.startswith("Heading 1") or name.startswith("Heading 2") or name.startswith("Heading 3"):
        return True
    if re.match(r"^제목\s*[123]", name):
        return True
    return False


# ── 표 변환 ──────────────────────────────────────────────────────────────────
def _cell_vmerge_state(cell) -> str:
    """vMerge 셀 상태 반환: 'restart' (시작) / 'continue' (계속) / '' (병합 없음).

    참고: 이 함수만으로는 python-docx의 cell.row.cells 동작을 완전히 추적할 수 없다.
    python-docx는 vMerge 그룹 전체에서 동일 _tc 인스턴스를 반환하므로, 호출자는
    _tc 객체 identity 비교(전/후 행과 동일하면 vMerge continue)도 함께 사용해야 한다.
    """
    try:
        tc = cell._tc
        for child in tc.tcPr.iterchildren() if tc.tcPr is not None else []:
            tag = child.tag
            if tag.endswith("}vMerge"):
                val = child.get(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"
                )
                return "restart" if val == "restart" else "continue"
    except Exception:
        pass
    return ""


def _escape_cell_text(text: str) -> str:
    """셀 텍스트를 파이프 마크다운 안전 형식으로 변환."""
    if not text:
        return ""
    # 개행 → 공백, 파이프 이스케이프, 컷
    cleaned = re.sub(r"\s*\n+\s*", " ", text).strip()
    cleaned = cleaned.replace("|", "&#124;")
    if len(cleaned) > _TABLE_CELL_MAX:
        cleaned = cleaned[: _TABLE_CELL_MAX - 1] + "…"
    return cleaned


def _table_to_markdown(table) -> tuple[str, list[dict]]:
    """표를 파이프 마크다운 + 병합 메타 정보로 변환.

    vMerge 감지 전략:
      1순위: cell._tc identity 비교 — python-docx는 vMerge 그룹 전체에 대해
             동일 _tc 인스턴스를 row.cells에 반환한다. 직전 행의 동일 컬럼
             _tc가 같으면 병합 continue.
      2순위: tcPr 내 w:vMerge 요소 검사 (보조).

    Returns:
        (markdown_string, merge_info_list)
        merge_info_list: [{row, col, merged_rows, value}, ...]
    """
    if not table.rows:
        return "", []

    rows_text: list[list[str]] = []
    merge_info: list[dict] = []

    n_cols = len(table.rows[0].cells) if table.rows else 0

    # vMerge 추적: 컬럼별 현재 병합 시작 행 + 시작 텍스트
    vmerge_start_row: dict[int, int] = {}
    vmerge_start_text: dict[int, str] = {}
    vmerge_run_len: dict[int, int] = {}
    prev_tc_per_col: dict[int, object] = {}     # 직전 행의 동일 컬럼 _tc

    for r_idx, row in enumerate(table.rows):
        row_cells: list[str] = []
        for c_idx, cell in enumerate(row.cells):
            tc = cell._tc
            xml_state = _cell_vmerge_state(cell)
            # 직전 행의 동일 컬럼과 _tc가 동일하면 vMerge 그룹 내부 (continue)
            is_continue = (
                xml_state == "continue"
                or (prev_tc_per_col.get(c_idx) is not None and prev_tc_per_col[c_idx] is tc)
            )
            raw_text = cell.text or ""

            if is_continue and c_idx in vmerge_run_len:
                row_cells.append("")  # 병합 계속 → 빈 칸
                vmerge_run_len[c_idx] += 1
            else:
                # 이전 병합 종료
                if c_idx in vmerge_run_len and vmerge_run_len[c_idx] > 1:
                    merge_info.append({
                        "row": vmerge_start_row[c_idx],
                        "col": c_idx,
                        "merged_rows": vmerge_run_len[c_idx],
                        "value": vmerge_start_text[c_idx],
                    })
                vmerge_start_row[c_idx]  = r_idx
                vmerge_start_text[c_idx] = raw_text.strip()
                vmerge_run_len[c_idx]    = 1
                row_cells.append(_escape_cell_text(raw_text))

            prev_tc_per_col[c_idx] = tc
        rows_text.append(row_cells)

    # 마지막 행까지 진행한 후 미종료 병합 flush
    for c_idx, run_len in vmerge_run_len.items():
        if run_len > 1:
            merge_info.append({
                "row": vmerge_start_row[c_idx],
                "col": c_idx,
                "merged_rows": run_len,
                "value": vmerge_start_text[c_idx],
            })

    if not rows_text or not n_cols:
        return "", merge_info

    # 첫 행 = 헤더
    header  = rows_text[0]
    body    = rows_text[1:]
    n_cols  = max(n_cols, max(len(r) for r in rows_text))

    def _pad(row: list[str]) -> list[str]:
        return row + [""] * (n_cols - len(row))

    lines: list[str] = []
    lines.append("| " + " | ".join(_pad(header)) + " |")
    lines.append("| " + " | ".join(["---"] * n_cols) + " |")
    for row in body:
        lines.append("| " + " | ".join(_pad(row)) + " |")

    return "\n".join(lines), merge_info


# ── 본문 순회 ─────────────────────────────────────────────────────────────────
def _iter_body_elements(doc) -> list[tuple[str, Any]]:
    """문서 본문을 순서대로 (kind, obj) 튜플 리스트로 반환.

    kind = "paragraph" | "table"
    obj  = Paragraph | Table
    """
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    body = doc.element.body
    p_tag = qn("w:p")
    tbl_tag = qn("w:tbl")

    para_by_elem = {p._element: p for p in doc.paragraphs}
    tbl_by_elem  = {t._element: t for t in doc.tables}

    items: list[tuple[str, Any]] = []
    for child in body.iterchildren():
        if child.tag == p_tag:
            obj = para_by_elem.get(child)
            if obj is not None:
                items.append(("paragraph", obj))
        elif child.tag == tbl_tag:
            obj = tbl_by_elem.get(child)
            if obj is not None:
                items.append(("table", obj))
    return items


# ── 분할 보조 ─────────────────────────────────────────────────────────────────
def _trigger_new_block(p, after_force_cut: bool = False) -> bool:
    """단락 p가 새 블록을 시작해야 하는지 판단 (R1~R4)."""
    text = (p.text or "").strip()
    if not text:
        return False
    # R1: Heading 스타일
    if _para_heading_style(p):
        return True
    # F1: TOC 항목은 R3 트리거 제외
    is_toc = bool(_RE_TOC.search(p.text or ""))
    # R4: 캡션 패턴
    if _RE_CAPTION.search(text):
        return True
    # R3: 섹션 번호 패턴 (TOC 제외)
    if not is_toc and _RE_SECTION_NUM.search(text):
        return True
    # R2: 폰트 visual
    if _para_visual_heading(p):
        return True
    return False


def _force_cut_needed(block: dict) -> bool:
    """F3: 단락 30개 또는 글자 8000자 초과 시 분할.

    block["paragraphs"]는 1차 raw_blocks 단계에서 Paragraph 객체 리스트.
    """
    paras = block["paragraphs"]
    if len(paras) >= _FORCE_CUT_PARAS:
        return True
    char_count = sum(len(p.text or "") for p in paras)
    return char_count >= _FORCE_CUT_CHARS


def _block_text_only_orphan(block: dict) -> bool:
    """F2: 단락 1개 + R3/캡션 헤더 + 표 없음 → 다음 블록에 흡수."""
    if block.get("table_markdown"):
        return False
    paras = block.get("paragraphs") or []
    if len(paras) != 1:
        return False
    text = paras[0].strip()
    if not text:
        return False
    if _RE_SECTION_NUM.search(text) or _RE_CAPTION.search(text):
        return True
    return False


def _is_toc_para(p) -> bool:
    """단락이 목차 항목인지 (탭 + 페이지번호 패턴)."""
    return bool(_RE_TOC.search(p.text or ""))


# ── page_header_text 폴백 ───────────────────────────────────────────────────
def _header_from_paragraph(p) -> str | None:
    """폴백 A: Heading 1/2/3 스타일 단락. 폴백 B: visual heading."""
    text = (p.text or "").strip()
    if not text:
        return None
    if _RE_FIGURE.match(text) or _RE_CONTINUE.search(text):
        return None
    return text[:_HEADER_FALLBACK_LEN]


def _decide_header(block_paragraphs: list, table_markdown: str | None,
                   block_num: int, breadcrumbs: list[str]) -> str:
    """page_header_text 폴백 A→B→C→D→E→F 순서."""
    # 폴백 A: 첫 Heading 스타일 단락
    for p in block_paragraphs:
        if _para_heading_style(p):
            h = _header_from_paragraph(p)
            if h:
                return _with_breadcrumbs(h, breadcrumbs)
    # 폴백 B: 첫 visual heading
    for p in block_paragraphs:
        if _para_visual_heading(p):
            h = _header_from_paragraph(p)
            if h:
                return _with_breadcrumbs(h, breadcrumbs)
    # 폴백 C: 캡션 패턴
    for p in block_paragraphs:
        text = (p.text or "").strip()
        m = _RE_CAPTION.search(text)
        if m:
            return _with_breadcrumbs(text[:_HEADER_FALLBACK_LEN], breadcrumbs)
    # 폴백 D: 표 첫 행 텍스트
    if table_markdown:
        first_line = table_markdown.split("\n", 1)[0]
        cells = [c.strip() for c in first_line.strip("|").split("|") if c.strip()]
        if cells:
            joined = " · ".join(cells)
            if joined:
                return _with_breadcrumbs(joined[:_HEADER_FALLBACK_LEN], breadcrumbs)
    # 폴백 E: 첫 비어있지 않은 단락
    for p in block_paragraphs:
        text = (p.text or "").strip()
        if text and not _RE_FIGURE.match(text) and not _RE_CONTINUE.search(text):
            return _with_breadcrumbs(text[:_HEADER_FALLBACK_LEN], breadcrumbs)
    # 폴백 F: 디폴트
    return _with_breadcrumbs(f"(블록 {block_num})", breadcrumbs)


def _with_breadcrumbs(header: str, breadcrumbs: list[str]) -> str:
    """breadcrumbs 가 있으면 'A > B > C > header' 형태로 누적."""
    if not breadcrumbs:
        return header
    chain = " > ".join(breadcrumbs + [header])
    return chain[:_HEADER_FALLBACK_LEN * 2]  # breadcrumbs는 길이 제한 완화


# ── 메인 함수 ─────────────────────────────────────────────────────────────────
def split_docx_to_blocks(path: str) -> list[dict]:
    """DOCX 파일을 의미 단위 블록으로 분할.

    각 블록 dict:
      {
        "block_num": int,
        "header_text": str,
        "paragraphs": [str],
        "table_markdown": str | None,
        "merge_info": [{"row", "col", "merged_rows", "value"}],
      }
    """
    from docx import Document
    doc = Document(path)
    items = _iter_body_elements(doc)

    # ── 1차: 원본 블록 분할 (R1~R5 + TOC 압축) ──────────────────────────────
    raw_blocks: list[dict] = []           # {paragraphs: [Paragraph], table: Table | None}
    current: dict | None = None
    toc_buffer: list = []                 # 연속된 TOC 단락 버퍼

    def _flush_toc():
        nonlocal current
        if not toc_buffer:
            return
        # F1: 연속 TOC 단락들은 "(목차)" 단일 블록으로 압축
        toc_block = {"paragraphs": list(toc_buffer), "table": None, "is_toc": True}
        raw_blocks.append(toc_block)
        toc_buffer.clear()
        current = None

    def _flush_current():
        nonlocal current
        if current and (current["paragraphs"] or current["table"] is not None):
            raw_blocks.append(current)
        current = None

    for kind, obj in items:
        if kind == "paragraph":
            text = (obj.text or "").strip()
            if not text:
                continue

            # TOC 항목 누적
            if _is_toc_para(obj):
                _flush_current()
                toc_buffer.append(obj)
                continue
            # 비-TOC 단락 등장 → 누적된 TOC flush
            if toc_buffer:
                _flush_toc()

            # R1~R4 트리거 검사
            if _trigger_new_block(obj):
                _flush_current()
                current = {"paragraphs": [obj], "table": None, "is_toc": False}
                # F3: force-cut 검사 (현 블록은 시작이라 단락 1개 → 트리거 안 됨)
            else:
                if current is None:
                    current = {"paragraphs": [obj], "table": None, "is_toc": False}
                else:
                    current["paragraphs"].append(obj)
                    # F3: force-cut → 새 블록 시작, "(계속)" 표시
                    if _force_cut_needed(current):
                        # 마지막 단락을 다음 블록으로 옮기지 않고 끊음
                        _flush_current()
                        current = {"paragraphs": [], "table": None, "is_toc": False,
                                   "force_cut_continue": True}

        elif kind == "table":
            # 표는 항상 단독 블록의 마지막 요소
            if toc_buffer:
                _flush_toc()
            if current is None or current["table"] is not None:
                _flush_current()
                current = {"paragraphs": [], "table": obj, "is_toc": False}
            else:
                current["table"] = obj
            _flush_current()  # 표 종료 후 블록 닫음

    if toc_buffer:
        _flush_toc()
    _flush_current()

    # ── 2차: F2 Forward merge orphan headers ──────────────────────────────────
    merged_blocks: list[dict] = []
    pending_breadcrumbs: list[str] = []
    i = 0
    while i < len(raw_blocks):
        blk = raw_blocks[i]
        para_texts = [(p.text or "").strip() for p in blk["paragraphs"]]

        # orphan: 단락 1개 + R3/캡션 매칭 + 표 없음
        is_orphan = (
            blk.get("table") is None
            and not blk.get("is_toc")
            and len(blk["paragraphs"]) == 1
        )
        if is_orphan:
            single_text = para_texts[0]
            if single_text and (_RE_SECTION_NUM.search(single_text)
                                or _RE_CAPTION.search(single_text)
                                or _para_heading_style(blk["paragraphs"][0])
                                or _para_visual_heading(blk["paragraphs"][0])):
                # 다음 블록에 흡수: breadcrumbs로 누적
                pending_breadcrumbs.append(single_text[:_HEADER_FALLBACK_LEN])
                i += 1
                continue

        # 일반 블록: breadcrumbs 부여 후 합류
        blk["breadcrumbs"] = list(pending_breadcrumbs)
        pending_breadcrumbs = []
        merged_blocks.append(blk)
        i += 1

    # 남은 pending_breadcrumbs는 마지막 dangling orphan(자체로 블록)으로 생성
    if pending_breadcrumbs:
        # 빈 블록 — 표 없음, 헤더만
        merged_blocks.append({
            "paragraphs": [],
            "table": None,
            "is_toc": False,
            "breadcrumbs": pending_breadcrumbs[:-1],
            "_orphan_header": pending_breadcrumbs[-1],
        })

    # ── 3차: dict 변환 (header_text + table_markdown + merge_info) ───────────
    result: list[dict] = []
    for idx, blk in enumerate(merged_blocks, start=1):
        para_objs = blk["paragraphs"]
        para_texts = [(p.text or "").strip() for p in para_objs if (p.text or "").strip()]
        table_md, merge_info = ("", []) if blk["table"] is None else _table_to_markdown(blk["table"])

        breadcrumbs = blk.get("breadcrumbs") or []

        if blk.get("is_toc"):
            header_text = _with_breadcrumbs("(목차)", breadcrumbs)
        elif blk.get("_orphan_header"):
            header_text = _with_breadcrumbs(blk["_orphan_header"], breadcrumbs)
        else:
            header_text = _decide_header(para_objs, table_md or None, idx, breadcrumbs)

        # force-cut 표시
        if blk.get("force_cut_continue"):
            header_text = f"{header_text} (계속)"

        result.append({
            "block_num":      idx,
            "header_text":    header_text,
            "paragraphs":     para_texts,
            "table_markdown": table_md or None,
            "merge_info":     merge_info,
        })

    return result


# ── source text 생성 (분류·추출 공통 입력) ──────────────────────────────────
def _detect_label_patterns(paragraphs: list[str]) -> list[str]:
    """단락 리스트에서 발견된 글머리 라벨 종류 반환 (LLM 힌트용)."""
    found: list[str] = []
    for name, pat in _LABEL_PATTERNS.items():
        if any(pat.match(p) for p in paragraphs):
            found.append(name)
    return found


def get_block_source_text(block: dict) -> str:
    """블록의 헤더 + 단락 + 표 미리보기를 결합한 텍스트.

    6,000자 초과 시 앞 4,000 + 뒤 2,000 (중간 "[...생략...]") 으로 컷.
    설계지침 글머리 라벨(1)/가)/① 등) 감지되면 헤더에 힌트 추가 — LLM 이
    design_guidelines_grouped[] 계층 보존 시 참조.
    """
    parts: list[str] = []
    parts.append(f"[HEADER] {block.get('header_text', '')}")

    paragraphs = block.get("paragraphs") or []
    labels_found = _detect_label_patterns(paragraphs) if paragraphs else []
    if labels_found:
        parts.append(f"[LABEL_PATTERNS_DETECTED] {' '.join(labels_found)}")
    if paragraphs:
        parts.append("[PARAGRAPHS]")
        parts.append("\n".join(paragraphs))

    tbl = block.get("table_markdown")
    if tbl:
        lines = tbl.split("\n")
        # 헤더 2줄(헤더 행 + 구분선) + 본문 최대 _TABLE_PREVIEW_ROWS 행
        head_lines = lines[:2]
        body_lines = lines[2:2 + _TABLE_PREVIEW_ROWS]
        if len(lines) > 2 + _TABLE_PREVIEW_ROWS:
            body_lines.append(f"| ... ({len(lines) - 2 - _TABLE_PREVIEW_ROWS}행 생략) |")
        parts.append("[TABLE]")
        parts.append("\n".join(head_lines + body_lines))

    text = "\n\n".join(parts)
    if len(text) <= _SOURCE_TEXT_CAP:
        return text
    return text[:_SOURCE_HEAD_CAP] + "\n\n[...생략...]\n\n" + text[-_SOURCE_TAIL_CAP:]
